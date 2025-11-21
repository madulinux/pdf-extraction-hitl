# Utility functions for document processing

import os
import re
import logging
from typing import Dict, List, Any, Optional

# Configure logging
logger = logging.getLogger("docservice.processor")

def suggest_variable_type(variable_name: str) -> str:
    """
    Suggest a variable type based on the variable name.
    
    Args:
        variable_name: Name of the variable
        
    Returns:
        Suggested variable type (string, number, date, image, etc.)
    """
    # Check for common patterns in variable names
    if any(keyword in variable_name.lower() for keyword in ['image', 'photo', 'picture', 'gambar', 'foto']):
        return 'image'
    
    if any(keyword in variable_name.lower() for keyword in ['date', 'tanggal', 'tgl', 'waktu', 'time', 'birth']):
        return 'date'
    
    if any(keyword in variable_name.lower() for keyword in ['price', 'harga', 'cost', 'biaya', 'amount', 'jumlah', 'total', 'sum', 'count']):
        return 'number'
    
    if any(keyword in variable_name.lower() for keyword in ['email', 'mail']):
        return 'email'
    
    if any(keyword in variable_name.lower() for keyword in ['phone', 'telp', 'telepon', 'hp', 'handphone', 'mobile']):
        return 'phone'
    
    if any(keyword in variable_name.lower() for keyword in ['address', 'alamat', 'location', 'lokasi']):
        return 'address'
    
    # Default to string
    return 'string'

def clean_variable_name(variable_name: str) -> str:
    """
    Clean a variable name by removing invalid characters.
    
    Args:
        variable_name: Name of the variable
        
    Returns:
        Cleaned variable name
    """
    # Remove special characters except underscore
    cleaned = re.sub(r'[^\w]', '_', variable_name)
    
    # Remove leading digits or underscores
    cleaned = re.sub(r'^[\d_]+', '', cleaned)
    
    # Convert to lowercase
    cleaned = cleaned.lower()
    
    # If the name is empty after cleaning, use a default name
    if not cleaned:
        cleaned = 'variable'
    
    return cleaned

def format_variable_value(value: Any, variable_type: str) -> str:
    """
    Format a variable value based on its type.
    
    Args:
        value: Variable value
        variable_type: Type of the variable
        
    Returns:
        Formatted variable value as a string
    """
    if value is None:
        return ''
    
    if variable_type == 'date':
        # Try to format as a date
        try:
            from datetime import datetime
            if isinstance(value, str):
                # Try common date formats
                formats = ['%Y-%m-%d', '%d-%m-%Y', '%Y/%m/%d', '%d/%m/%Y', '%d %B %Y', '%B %d, %Y']
                for fmt in formats:
                    try:
                        date_obj = datetime.strptime(value, fmt)
                        return date_obj.strftime('%d %B %Y')  # Format as day month year
                    except ValueError:
                        continue
                # If no format matched, return as is
                return value
            elif isinstance(value, datetime):
                return value.strftime('%d %B %Y')  # Format as day month year
            else:
                return str(value)
        except Exception as e:
            logger.warning(f"Error formatting date value: {str(e)}")
            return str(value)
    
    if variable_type == 'number':
        # Try to format as a number
        try:
            if isinstance(value, (int, float)):
                return '{:,}'.format(value)  # Format with thousand separators
            else:
                return str(value)
        except Exception as e:
            logger.warning(f"Error formatting number value: {str(e)}")
            return str(value)
    
    # Default to string
    return str(value)

def get_docx_metadata(docx_path: str) -> Dict[str, Any]:
    """
    Get metadata from a DOCX file.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Dictionary of metadata
    """
    try:
        import docx
        
        # Load the DOCX file
        doc = docx.Document(docx_path)
        
        # Get core properties
        core_props = doc.core_properties
        
        # Create metadata dictionary
        metadata = {
            'author': core_props.author,
            'title': core_props.title,
            'subject': core_props.subject,
            'keywords': core_props.keywords,
            'created': core_props.created,
            'modified': core_props.modified,
            'last_modified_by': core_props.last_modified_by,
            'revision': core_props.revision,
            'category': core_props.category,
            'comments': core_props.comments,
            'content_status': core_props.content_status,
            'identifier': core_props.identifier,
            'language': core_props.language,
            'version': core_props.version
        }
        
        # Add document statistics
        metadata['paragraphs_count'] = len(doc.paragraphs)
        metadata['tables_count'] = len(doc.tables)
        metadata['sections_count'] = len(doc.sections)
        
        return metadata
    except Exception as e:
        logger.error(f"Error getting DOCX metadata: {str(e)}")
        return {}
