# Template processing module

import os
import re
import zipfile
import shutil
import tempfile
import logging
from typing import Dict, List, Any, Optional, Tuple
from lxml import etree

# Configure logging
logger = logging.getLogger("docservice.processor")

def process_template(template_path: str, output_path: str, variables: Dict[str, Any]) -> str:
    """
    Process a DOCX template file by replacing variables with values.
    
    Args:
        template_path: Path to the DOCX template file
        output_path: Path to save the processed DOCX file
        variables: Dictionary of variable names and values
        
    Returns:
        Path to the processed DOCX file
    """
    # Check if file exists
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found: {template_path}")
        
    # Log the variables being injected
    # logger.info(f"Injecting variables: {variables}")
    
    # Create a copy of the template file
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    shutil.copy2(template_path, output_path)
    
    # Process the template with direct XML replacement (more reliable for complex templates)
    try:
        # Create a temporary directory to extract the DOCX file
        temp_dir = tempfile.mkdtemp(prefix="temp_docx_", dir=os.path.dirname(output_path))
        
        # Extract the DOCX file to the temporary directory
        with zipfile.ZipFile(output_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
        
        # Process all XML files in the DOCX file
        xml_files = [
            os.path.join(temp_dir, 'word', 'document.xml'),
            os.path.join(temp_dir, 'word', 'header1.xml'),
            os.path.join(temp_dir, 'word', 'header2.xml'),
            os.path.join(temp_dir, 'word', 'header3.xml'),
            os.path.join(temp_dir, 'word', 'footer1.xml'),
            os.path.join(temp_dir, 'word', 'footer2.xml'),
            os.path.join(temp_dir, 'word', 'footer3.xml')
        ]
        
        # Process each XML file if it exists
        for xml_file in xml_files:
            if os.path.exists(xml_file):
                _process_xml_file(xml_file, variables)
        
        # Create a new ZIP file with the modified content
        new_docx = output_path + ".new"
        _create_docx_from_directory(temp_dir, new_docx)
        
        # Replace the original file with the new one
        shutil.move(new_docx, output_path)
        
        # Clean up the temporary directory
        shutil.rmtree(temp_dir)
        
        # logger.info("Direct XML replacement completed")
        
        # Verify that all variables have been replaced
        _verify_output_document(output_path, variables)
        
        return output_path
    except Exception as e:
        logger.error(f"Error processing template: {str(e)}")
        raise

def _process_xml_file(xml_file: str, variables: Dict[str, Any]) -> None:
    """
    Process a single XML file in the DOCX package.
    
    Args:
        xml_file: Path to the XML file
        variables: Dictionary of variable names and values
    """
    try:
        # First try with lxml for precise XML manipulation
        try:
            # Parse the XML file
            parser = etree.XMLParser(remove_blank_text=True)
            tree = etree.parse(xml_file, parser)
            root = tree.getroot()
            
            # Find all text elements
            namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            text_elements = root.xpath('//w:t', namespaces=namespaces)
            
            # First pass: Try to replace variables in individual text elements
            for element in text_elements:
                if element.text:
                    # Replace variables with single braces: {variable}
                    for var_name, var_value in variables.items():
                        if '{' + var_name + '}' in element.text:
                            element.text = element.text.replace('{' + var_name + '}', str(var_value))
                    
                    # Replace variables with double braces: {{variable}}
                    for var_name, var_value in variables.items():
                        if '{{' + var_name + '}}' in element.text:
                            element.text = element.text.replace('{{' + var_name + '}}', str(var_value))
            
            # Second pass: Handle fragmented variables across multiple text elements
            # Get all paragraphs
            paragraphs = root.xpath('//w:p', namespaces=namespaces)
            
            for paragraph in paragraphs:
                # Get all text elements in this paragraph
                p_text_elements = paragraph.xpath('.//w:t', namespaces=namespaces)
                
                # Skip if there are fewer than 2 text elements (not fragmented)
                if len(p_text_elements) < 2:
                    continue
                
                # Reconstruct the paragraph text
                paragraph_text = ''.join([el.text if el.text else '' for el in p_text_elements])
                
                # Check if this paragraph contains any variables
                contains_variable = False
                for var_name in variables.keys():
                    if '{' + var_name + '}' in paragraph_text or '{{' + var_name + '}}' in paragraph_text:
                        contains_variable = True
                        break
                
                if not contains_variable:
                    continue
                
                # Find potential variable fragments
                opening_brace_elements = []
                variable_name_elements = []
                closing_brace_elements = []
                
                # Identify elements that might contain parts of variables
                for i, el in enumerate(p_text_elements):
                    if el.text:
                        if '{' in el.text:
                            opening_brace_elements.append((i, el, '{'))
                        if '{{' in el.text:
                            opening_brace_elements.append((i, el, '{{'))
                        
                        if '}' in el.text:
                            closing_brace_elements.append((i, el, '}'))
                        if '}}' in el.text:
                            closing_brace_elements.append((i, el, '}}'))
                        
                        # Check if this element might contain a variable name
                        for var_name in variables.keys():
                            if var_name in el.text and not ('{' in el.text or '}' in el.text):
                                variable_name_elements.append((i, el, var_name))
                
                # Process potential variable fragments
                for open_idx, open_el, open_brace in opening_brace_elements:
                    for close_idx, close_el, close_brace in closing_brace_elements:
                        if open_idx < close_idx:  # Only process if opening comes before closing
                            # Make sure braces match (single or double)
                            if (open_brace == '{' and close_brace == '}') or (open_brace == '{{' and close_brace == '}}'):
                                # Extract the text between opening and closing braces
                                fragment = ''.join([p_text_elements[i].text if p_text_elements[i].text else '' 
                                                  for i in range(open_idx, close_idx + 1)])
                                
                                # Check if this fragment contains any variable
                                for var_name, var_value in variables.items():
                                    var_pattern_single = '{' + var_name + '}'
                                    var_pattern_double = '{{' + var_name + '}}'
                                    
                                    if var_pattern_single in fragment:
                                        # Found a fragmented variable! Now replace it
                                        # First, handle the opening brace element
                                        if '{' in open_el.text:
                                            # Find where the opening brace is in the text
                                            brace_pos = open_el.text.find('{')
                                            
                                            # Replace just the opening part
                                            prefix = open_el.text[:brace_pos]
                                            open_el.text = prefix + str(var_value)
                                            
                                            # Clear all elements in between
                                            for i in range(open_idx + 1, close_idx):
                                                p_text_elements[i].text = ''
                                            
                                            # Clear the closing part
                                            if '}' in close_el.text:
                                                suffix_pos = close_el.text.find('}') + 1
                                                close_el.text = close_el.text[suffix_pos:]
                                    
                                    elif var_pattern_double in fragment:
                                        # Found a fragmented variable! Now replace it
                                        # First, handle the opening brace element
                                        if '{{' in open_el.text:
                                            # Find where the opening brace is in the text
                                            brace_pos = open_el.text.find('{{')
                                            
                                            # Replace just the opening part
                                            prefix = open_el.text[:brace_pos]
                                            open_el.text = prefix + str(var_value)
                                            
                                            # Clear all elements in between
                                            for i in range(open_idx + 1, close_idx):
                                                p_text_elements[i].text = ''
                                            
                                            # Clear the closing part
                                            if '}}' in close_el.text:
                                                suffix_pos = close_el.text.find('}}') + 2
                                                close_el.text = close_el.text[suffix_pos:]
            
            # Write the modified XML back to the file
            tree.write(xml_file, encoding='UTF-8', xml_declaration=True)
            return  # Success, no need for fallback
        except Exception as xml_error:
            logger.warning(f"XML parsing approach failed: {str(xml_error)}. Trying fallback...")
            # Continue to fallback method
        
        # Fallback: Try with direct string replacement
        _process_xml_file_fallback(xml_file, variables)
    except Exception as e:
        logger.error(f"All methods failed processing XML file {xml_file}: {str(e)}")
        # Try one last desperate approach with direct file editing
        try:
            with open(xml_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Replace all variables in the content
            for var_name, var_value in variables.items():
                # Replace with single braces
                content = content.replace('{' + var_name + '}', str(var_value))
                # Replace with double braces
                content = content.replace('{{' + var_name + '}}', str(var_value))
            
            # Write the content back
            with open(xml_file, 'w', encoding='utf-8') as f:
                f.write(content)
            
        except Exception as last_error:
            logger.error(f"Even last-resort method failed for {xml_file}: {str(last_error)}")

def _process_xml_file_fallback(xml_file: str, variables: Dict[str, Any]) -> None:
    """
    Fallback method to process XML file using direct string replacement.
    
    Args:
        xml_file: Path to the XML file
        variables: Dictionary of variable names and values
    """
    try:
        # Read the XML file as text
        with open(xml_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # First pass: Replace simple variables
        for var_name, var_value in variables.items():
            # Replace variables with single braces: {variable}
            content = content.replace('{' + var_name + '}', str(var_value))
            # Replace variables with double braces: {{variable}}
            content = content.replace('{{' + var_name + '}}', str(var_value))
        
        # Second pass: Handle fragmented variables
        # Look for opening braces followed by variable name
        for var_name, var_value in variables.items():
            # Pattern for single braces split across tags
            pattern_single = r'\{\s*</w:t>.*?<w:t[^>]*>\s*' + re.escape(var_name) + r'\s*</w:t>.*?<w:t[^>]*>\s*\}'
            content = re.sub(pattern_single, str(var_value), content)
            
            # Pattern for double braces split across tags
            pattern_double = r'\{\{\s*</w:t>.*?<w:t[^>]*>\s*' + re.escape(var_name) + r'\s*</w:t>.*?<w:t[^>]*>\s*\}\}'
            content = re.sub(pattern_double, str(var_value), content)
            
            # Pattern for variable name split across tags
            var_parts = var_name.split('_')
            if len(var_parts) > 1:
                for i in range(1, len(var_parts)):
                    prefix = '_'.join(var_parts[:i])
                    suffix = '_'.join(var_parts[i:])
                    pattern = r'\{\s*' + re.escape(prefix) + r'\s*</w:t>.*?<w:t[^>]*>\s*' + re.escape('_' + suffix) + r'\s*\}'
                    content = re.sub(pattern, str(var_value), content)
                    
                    # Same for double braces
                    pattern = r'\{\{\s*' + re.escape(prefix) + r'\s*</w:t>.*?<w:t[^>]*>\s*' + re.escape('_' + suffix) + r'\s*\}\}'
                    content = re.sub(pattern, str(var_value), content)
        
        # Third pass: Clean up any leftover variable fragments
        # Remove any leftover opening braces
        content = re.sub(r'\{[a-zA-Z0-9_]*</w:t>', '</w:t>', content)
        content = re.sub(r'\{\{[a-zA-Z0-9_]*</w:t>', '</w:t>', content)
        
        # Remove any leftover closing braces
        content = re.sub(r'<w:t[^>]*>[a-zA-Z0-9_]*\}', '<w:t>', content)
        content = re.sub(r'<w:t[^>]*>[a-zA-Z0-9_]*\}\}', '<w:t>', content)
        
        # Write the modified content back
        with open(xml_file, 'w', encoding='utf-8') as f:
            f.write(content)
            
    except Exception as e:
        logger.error(f"Error in fallback processing for {xml_file}: {str(e)}")
        raise

def _create_docx_from_directory(directory: str, output_path: str) -> None:
    """
    Create a DOCX file from a directory.
    
    Args:
        directory: Path to the directory containing the DOCX contents
        output_path: Path to save the DOCX file
    """
    try:
        # Create a new ZIP file
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # Walk through the directory and add all files to the ZIP
            for root, dirs, files in os.walk(directory):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, directory)
                    zip_file.write(file_path, arcname)
    except Exception as e:
        logger.error(f"Error creating DOCX from directory: {str(e)}")
        raise

def _verify_output_document(output_path: str, variables: Dict[str, Any]) -> None:
    """
    Verify that all variables have been replaced in the output document.
    
    Args:
        output_path: Path to the output DOCX file
        variables: Dictionary of variable names and values
    """
    try:
        # Extract the DOCX file to a temporary directory
        temp_dir = tempfile.mkdtemp(prefix="verify_docx_", dir=os.path.dirname(output_path))
        
        with zipfile.ZipFile(output_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)
        
        # Check all XML files for unreplaced variables
        xml_files = [
            os.path.join(temp_dir, 'word', 'document.xml'),
            os.path.join(temp_dir, 'word', 'header1.xml'),
            os.path.join(temp_dir, 'word', 'header2.xml'),
            os.path.join(temp_dir, 'word', 'header3.xml'),
            os.path.join(temp_dir, 'word', 'footer1.xml'),
            os.path.join(temp_dir, 'word', 'footer2.xml'),
            os.path.join(temp_dir, 'word', 'footer3.xml')
        ]
        
        unreplaced_variables = set()
        
        for xml_file in xml_files:
            if os.path.exists(xml_file):
                with open(xml_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                    
                    # Check for unreplaced variables
                    for var_name in variables.keys():
                        if '{' + var_name + '}' in content or '{{' + var_name + '}}' in content:
                            unreplaced_variables.add(var_name)
        
        # Clean up the temporary directory
        shutil.rmtree(temp_dir)
        
        # Log unreplaced variables
        if unreplaced_variables:
            logger.warning(f"Found unreplaced variables in output document: {list(unreplaced_variables)}")
        
    except Exception as e:
        logger.error(f"Error verifying output document: {str(e)}")

def process_image_variables(template_path: str, output_path: str, image_variables: Dict[str, str]) -> str:
    """
    Process image variables in a DOCX template.
    
    Args:
        template_path: Path to the DOCX template file
        output_path: Path to save the processed DOCX file
        image_variables: Dictionary of image variable names and image file paths
        
    Returns:
        Path to the processed DOCX file
    """
    try:
        # Create a copy of the template file
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        shutil.copy2(template_path, output_path)
        
        # Process the template with python-docx
        import docx
        doc = docx.Document(output_path)
        
        # Process each paragraph
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                for var_name, image_path in image_variables.items():
                    if '{' + var_name + '}' in run.text or '{{' + var_name + '}}' in run.text:
                        # Clear the run text
                        run.text = ''
                        # Add the image
                        run.add_picture(image_path)
        
        # Process each table
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            for var_name, image_path in image_variables.items():
                                if '{' + var_name + '}' in run.text or '{{' + var_name + '}}' in run.text:
                                    # Clear the run text
                                    run.text = ''
                                    # Add the image
                                    run.add_picture(image_path)
        
        # Save the document
        doc.save(output_path)
        
        return output_path
    except Exception as e:
        logger.error(f"Error processing image variables: {str(e)}")
        raise
