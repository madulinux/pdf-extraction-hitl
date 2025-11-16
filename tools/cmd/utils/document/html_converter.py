# HTML conversion module

import os
import re
import zipfile
import logging
from typing import Dict, List, Any, Optional

# Configure logging
logger = logging.getLogger("docservice.processor")

def docx_to_html(docx_path: str, output_path: str = None) -> str:
    """
    Convert a DOCX file to HTML.
    
    Args:
        docx_path: Path to the DOCX file
        output_path: Path to save the HTML file (optional)
        
    Returns:
        HTML content as a string if output_path is None, otherwise path to the HTML file
    """
    # Check if file exists
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")
    
    # Generate output path if not provided
    if output_path is None:
        output_path = os.path.splitext(docx_path)[0] + ".html"
    
    # Create output directory if it doesn't exist
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # Try multiple conversion methods
    html_content = None
    
    # Method 1: Try with python-docx-html
    try:
        html_content = _convert_with_docx_html(docx_path)
    except Exception as e:
        logger.warning(f"python-docx-html conversion failed: {str(e)}")
    
    # Method 2: Try with mammoth
    if not html_content:
        try:
            html_content = _convert_with_mammoth(docx_path)
        except Exception as e:
            logger.warning(f"mammoth conversion failed: {str(e)}")
    
    # Method 3: Try with custom conversion
    if not html_content:
        try:
            html_content = _convert_with_custom(docx_path)
        except Exception as e:
            logger.warning(f"Custom conversion failed: {str(e)}")
    
    # If all methods failed, raise an exception
    if not html_content:
        raise Exception("All HTML conversion methods failed")
    
    # Write HTML content to file if output_path is provided
    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        # logger.info(f"Successfully converted DOCX to HTML: {output_path}")
        return output_path
    
    return html_content

def _convert_with_docx_html(docx_path: str) -> Optional[str]:
    """
    Convert a DOCX file to HTML using python-docx-html.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        HTML content as a string or None if conversion failed
    """
    try:
        import docx
        from docx_html import docx_to_html
        
        # Load the DOCX file
        doc = docx.Document(docx_path)
        
        # Convert to HTML
        html_content = docx_to_html(doc)
        
        # Check if the HTML content is valid
        if html_content and "<html" in html_content:
            # logger.info("Successfully converted DOCX to HTML using python-docx-html")
            return html_content
        else:
            logger.warning("python-docx-html returned invalid HTML content")
            return None
    except ImportError:
        logger.warning("python-docx-html module not installed")
        return None
    except Exception as e:
        logger.error(f"Error during python-docx-html conversion: {str(e)}")
        return None

def _convert_with_mammoth(docx_path: str) -> Optional[str]:
    """
    Convert a DOCX file to HTML using mammoth.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        HTML content as a string or None if conversion failed
    """
    try:
        import mammoth
        
        # Convert to HTML
        with open(docx_path, 'rb') as f:
            result = mammoth.convert_to_html(f)
            html_content = result.value
        
        # Check if the HTML content is valid
        if html_content and "<p" in html_content:
            # Wrap in HTML structure if needed
            if "<html" not in html_content:
                html_content = f"""<!DOCTYPE html>
                <html>
                <head>
                    <meta charset="UTF-8">
                    <title>Document</title>
                    <style>
                        body {{ font-family: Arial, sans-serif; }}
                        table {{ border-collapse: collapse; width: 100%; }}
                        th, td {{ border: 1px solid #ddd; padding: 8px; }}
                        th {{ background-color: #f2f2f2; }}
                    </style>
                </head>
                <body>
                {html_content}
                </body>
                </html>"""
            
            # logger.info("Successfully converted DOCX to HTML using mammoth")
            return html_content
        else:
            logger.warning("mammoth returned invalid HTML content")
            return None
    except ImportError:
        logger.warning("mammoth module not installed")
        return None
    except Exception as e:
        logger.error(f"Error during mammoth conversion: {str(e)}")
        return None

def _convert_with_custom(docx_path: str) -> Optional[str]:
    """
    Convert a DOCX file to HTML using a custom method.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        HTML content as a string or None if conversion failed
    """
    try:
        import docx
        
        # Load the DOCX file
        doc = docx.Document(docx_path)
        
        # Create HTML structure
        html = [
            '<!DOCTYPE html>',
            '<html>',
            '<head>',
            '    <meta charset="UTF-8">',
            '    <title>Document</title>',
            '    <style>',
            '        body { font-family: Arial, sans-serif; }',
            '        table { border-collapse: collapse; width: 100%; }',
            '        th, td { border: 1px solid #ddd; padding: 8px; }',
            '        th { background-color: #f2f2f2; }',
            '    </style>',
            '</head>',
            '<body>'
        ]
        
        # Process paragraphs
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
                
            # Check for heading style
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.replace('Heading', ''))
                html.append(f'<h{level}>{para.text}</h{level}>')
            else:
                html.append(f'<p>{para.text}</p>')
        
        # Process tables
        for table in doc.tables:
            html.append('<table>')
            
            # Process rows
            for i, row in enumerate(table.rows):
                html.append('<tr>')
                
                # Process cells
                for cell in row.cells:
                    # Use th for header row
                    if i == 0:
                        html.append(f'<th>{cell.text}</th>')
                    else:
                        html.append(f'<td>{cell.text}</td>')
                
                html.append('</tr>')
            
            html.append('</table>')
        
        # Close HTML structure
        html.extend(['</body>', '</html>'])
        
        return '\n'.join(html)
    except Exception as e:
        logger.error(f"Error converting DOCX to HTML: {str(e)}")
        raise
