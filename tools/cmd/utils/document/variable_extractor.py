# Variable extraction module

import os
import re
import zipfile
import logging
from typing import Dict, List, Any, Optional

# Configure logging
logger = logging.getLogger("docservice.processor")

def extract_variables_from_template(template_path: str) -> List[Dict[str, Any]]:
    """
    Extract variables from a DOCX template file.
    
    Args:
        template_path: Path to the DOCX template file
        
    Returns:
        List of dictionaries containing variable information
    """
    # Check if file exists
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template file not found: {template_path}")
        
    # Check if file is empty
    if os.path.getsize(template_path) == 0:
        raise ValueError(f"Template file is empty: {template_path}")
        
    # Check if file is too large
    max_size = 10 * 1024 * 1024  # 10MB
    if os.path.getsize(template_path) > max_size:
        raise ValueError(f"Template file is too large: {template_path}. Maximum size is 10MB.")
    
    # Try direct XML extraction first (most reliable for simple templates)
    try:
        # Check if it's a valid ZIP file (DOCX is a ZIP file)
        if not zipfile.is_zipfile(template_path):
            raise ValueError(f"The file is not a valid DOCX (ZIP) file: {template_path}")
            
        # Set up debug file if debugging is enabled
        debug_enabled = os.environ.get('DEBUG', 'false').lower() == 'true'
        debug_file = None
        
        if debug_enabled:
            debug_dir = 'debug'
            os.makedirs(debug_dir, exist_ok=True)
            debug_file = os.path.join(debug_dir, f"document_{os.path.basename(template_path)}.xml")
            # logger.info(f"Saving document XML for debugging to {debug_file}")
        
        # Open the DOCX file as a ZIP file
        with zipfile.ZipFile(template_path, 'r') as zip_ref:
            # List all files in the DOCX archive
            file_list = zip_ref.namelist()
            # logger.info(f"Files in DOCX archive: {file_list}")
            
            # Check if document.xml exists
            document_path = 'word/document.xml'
            if document_path not in file_list:
                raise ValueError(f"The DOCX file does not contain {document_path}")
                
            # Extract and read document.xml
            with zip_ref.open(document_path) as doc_file:
                content = doc_file.read().decode('utf-8')
                
                # Save XML content for debugging if enabled
                if debug_enabled:
                    with open(debug_file, 'w', encoding='utf-8') as f:
                        f.write(content)
                
                # Log a sample of the content for debugging
                content_sample = content[:500] + '...' if len(content) > 500 else content
                # logger.info(f"XML content sample: {content_sample}")
                
                # Preprocessing: Clean up the XML to make variable extraction easier
                # Remove all XML tags to get plain text content
                import re
                # First, extract all text content from w:t tags
                text_content = []
                for match in re.finditer(r'<w:t[^>]*>(.*?)</w:t>', content, re.DOTALL):
                    text_content.append(match.group(1))
                
                # Join all text content
                plain_text = ''.join(text_content)
                # logger.info(f"Extracted plain text (sample): {plain_text[:200]}...")
                
                # Find all variables in plain text
                var_pattern = r'\{([a-zA-Z0-9_\-\.]+)\}'
                variables = {}
                
                for match in re.finditer(var_pattern, plain_text):
                    var_name = match.group(1)
                    
                    # Skip Jinja2 control structures
                    if var_name.startswith(('if', 'for', 'end')):
                        continue
                        
                    if var_name in variables:
                        variables[var_name] += 1
                    else:
                        variables[var_name] = 1
                
                # Also check for {{variable}} format (double braces)
                double_brace_pattern = r'\{\{([a-zA-Z0-9_\-\.]+)\}\}'
                for match in re.finditer(double_brace_pattern, plain_text):
                    var_name = match.group(1)
                    
                    # Skip Jinja2 control structures
                    if var_name.startswith(('if', 'for', 'end')):
                        continue
                        
                    if var_name in variables:
                        variables[var_name] += 1
                    else:
                        variables[var_name] = 1
                
                # If still not finding variables, try an alternative approach
                if len(variables) < 2:
                    # logger.info("Still few variables found, trying alternative approach")
                    
                    # Try to find variables that might be split across XML tags
                    # First, simplify the XML by removing attributes but keeping tags
                    simplified_xml = re.sub(r'<([^>]*)\s+[^>]*>', r'<\1>', content)
                    
                    # Now look for text between <w:t> tags that contains { or }
                    fragments = []
                    for match in re.finditer(r'<w:t>([^<]*[{}][^<]*)</w:t>', simplified_xml):
                        fragments.append(match.group(1))
                    
                    # Join fragments and look for variables
                    joined_text = ''.join(fragments)
                    # logger.info(f"Joined text fragments (sample): {joined_text[:200]}...")
                    
                    # Look for variables in the joined text
                    for match in re.finditer(var_pattern, joined_text):
                        var_name = match.group(1)
                        
                        # Skip Jinja2 control structures
                        if var_name.startswith(('if', 'for', 'end')):
                            continue
                            
                        if var_name in variables:
                            variables[var_name] += 1
                        else:
                            variables[var_name] = 1
                    
                    # Also check for {{variable}} format in joined text
                    for match in re.finditer(double_brace_pattern, joined_text):
                        var_name = match.group(1)
                        
                        # Skip Jinja2 control structures
                        if var_name.startswith(('if', 'for', 'end')):
                            continue
                            
                        if var_name in variables:
                            variables[var_name] += 1
                        else:
                            variables[var_name] = 1
                
                # Last resort: Try to find variable names directly from the HTML output
                if len(variables) < 2 and debug_enabled:
                    # logger.info("Trying to extract variables from HTML output")
                    
                    # Check if we have an HTML output file to analyze
                    output_files = [f for f in os.listdir('storage/output') if f.endswith('.html')]
                    if output_files:
                        latest_html = os.path.join('storage/output', output_files[-1])
                        # logger.info(f"Analyzing HTML file: {latest_html}")
                        
                        try:
                            with open(latest_html, 'r', encoding='utf-8') as html_file:
                                html_content = html_file.read()
                                
                                # Find all variables in HTML content
                                for match in re.finditer(r'\{([a-zA-Z0-9_\-\.]+)\}', html_content):
                                    var_name = match.group(1)
                                    
                                    # Skip Jinja2 control structures
                                    if var_name.startswith(('if', 'for', 'end')):
                                        continue
                                        
                                    if var_name in variables:
                                        variables[var_name] += 1
                                    else:
                                        variables[var_name] = 1
                        except Exception as html_error:
                            logger.error(f"Error analyzing HTML: {str(html_error)}")
        
        # If we found variables with direct approach, return them
        if variables:
            # logger.info(f"Successfully extracted {len(variables)} variables using direct XML approach")
            
            # Convert to list of dictionaries with additional metadata
            from .utils import suggest_variable_type
            
            result = []
            for var_name, occurrences in variables.items():
                # Determine suggested type based on variable name
                suggested_type = suggest_variable_type(var_name)
                
                result.append({
                    "name": var_name,
                    "occurrences": occurrences,
                    "suggested_type": suggested_type
                })
            return result
            
        # If no variables found with direct approach, continue with DocxTemplate
        # logger.info("No variables found with direct approach, trying DocxTemplate...")
        
    except Exception as direct_error:
        # Log the error but continue with DocxTemplate approach
        logger.warning(f"Direct XML extraction failed: {str(direct_error)}. Trying DocxTemplate approach...")
    
    # Try the DocxTemplate approach
    try:
        # Try to open the file with python-docx first to verify it's a valid DOCX
        import docx
        try:
            # This will validate if it's a proper DOCX file
            doc_check = docx.Document(template_path)
            # Log some info about the document for diagnostics
            paragraphs_count = len(doc_check.paragraphs)
            tables_count = len(doc_check.tables)
            logger.info(f"DOCX validation: {paragraphs_count} paragraphs, {tables_count} tables")
        except Exception as docx_error:
            logger.error(f"Failed to validate DOCX with python-docx: {str(docx_error)}")
            raise ValueError(f"The file does not appear to be a valid DOCX document. Error: {str(docx_error)}")
        
        # If we get here, the file is at least a valid DOCX structure
        # Now try with DocxTemplate
        from docxtpl import DocxTemplate
        doc = DocxTemplate(template_path)
        
        # Verify the document was loaded properly
        if doc is None:
            raise ValueError(f"DocxTemplate returned None for: {template_path}")
            
        if not hasattr(doc, 'docx'):
            raise ValueError(f"DocxTemplate object has no 'docx' attribute: {template_path}")
            
        if doc.docx is None:
            raise ValueError(f"DocxTemplate.docx is None: {template_path}")
            
        # Get the document XML with error handling
        try:
            xml_content = doc.get_xml()
            if not xml_content:
                raise ValueError(f"Document loaded but XML content is empty")
                
            # Find all variables in the format {variable_name}
            pattern = r'\{([a-zA-Z0-9_\.]+)\}'
            variables = {}
            
            for match in re.finditer(pattern, xml_content):
                var_name = match.group(1)
                
                # Skip Jinja2 control structures
                if var_name.startswith(('if', 'for', 'end')):
                    continue
                    
                if var_name in variables:
                    variables[var_name] += 1
                else:
                    variables[var_name] = 1
            
            # Convert to list of dictionaries with additional metadata
            from .utils import suggest_variable_type
            
            result = []
            for var_name, occurrences in variables.items():
                # Determine suggested type based on variable name
                suggested_type = suggest_variable_type(var_name)
                
                result.append({
                    "name": var_name,
                    "occurrences": occurrences,
                    "suggested_type": suggested_type
                })
            
            # If no variables found, provide a helpful message
            if not result:
                logger.warning(f"No variables found in template: {template_path}")
                
            return result
        except Exception as xml_error:
            logger.error(f"Error extracting XML from DocxTemplate: {str(xml_error)}")
            raise ValueError(f"Failed to extract variables from template: {str(xml_error)}")
    except Exception as e:
        logger.error(f"Error extracting variables from template: {str(e)}")
        raise
